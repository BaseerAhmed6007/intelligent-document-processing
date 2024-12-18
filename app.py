import streamlit as st
import os
from io import BytesIO  # To handle file upload
from azure.core.credentials import AzureKeyCredential
from azure.ai.documentintelligence import DocumentIntelligenceClient
from azure.ai.textanalytics import TextAnalyticsClient
from docx import Document
from sklearn.metrics.pairwise import cosine_similarity
from sentence_transformers import SentenceTransformer
from azure.ai.language.conversations import ConversationAnalysisClient
from openai import AzureOpenAI
import numpy as np
import re
import cv2

# Helper functions and API clients here...
# Fetch secret keys from secret storage
azure_api_key = st.secrets['AZURE_API_KEY']
azure_endpoint = st.secrets['AZURE_ENDPOINT']
# azure_openai_endpoint = st.secrets['AZURE_OPENAI_ENDPOINT']  # Assuming you store this
# azure_openai_key = st.secrets['AZURE_OPENAI_KEY']  # Assuming you store this
text_analytics_api_key = st.secrets['TEXT_ANALYTICS_API_KEY']  # Assuming you store this
text_analytics_endpoint = st.secrets['TEXT_ANALYTICS_ENDPOINT']  # Assuming you store this
convers_analysis_api_key = st.secrets['CONVERSATION_ANALYSIS_API_KEY']  # Assuming you store this
convers_analysis_endpoint = st.secrets['CONVERSATION_ANALYSIS_ENDPOINT']  # Assuming you store this

# Initialize Azure Text Analytics client
text_analytics_client = TextAnalyticsClient(
    endpoint=text_analytics_endpoint,
    credential=AzureKeyCredential(text_analytics_api_key)
)
conversation_analysis_client = ConversationAnalysisClient(
    convers_analysis_endpoint, AzureKeyCredential(convers_analysis_api_key)
)

# Helper function to get words within a line's spans
def get_words(page, line):
    result = []
    for word in page.words:
        if _in_span(word, line.spans):
            result.append(word)
    return result

# Helper function to check if a word is within any of the spans
def _in_span(word, spans):
    for span in spans:
        if word.span.offset >= span.offset and (
            word.span.offset + word.span.length
        ) <= (span.offset + span.length):
            return True
    return False

def recognize_intent(user_command):
    if not user_command or not user_command.strip():
        raise ValueError("The input command is invalid. Please provide a non-empty command.")

    response = conversation_analysis_client.analyze_conversation(
        task={
            "kind": "Conversation",
            "analysisInput": {
                "conversationItem": {
                    "id": "1",
                    "text": user_command,
                    "participantId": "1"
                }
            },
            "parameters": {
                "projectName": "ConvUnder",
                "deploymentName": "Conversationn"
            }
        }
    )

    intents = response["result"]["prediction"]["topIntent"]
    return intents

def process_intent(intent, text, openai_client):
    if intent == "summary":
        return summarize_text(text, openai_client)
    elif intent == "RedactPII":
        return redact_pii(text)
    elif intent == "GetEntities":
        return extract_entities(text)
    elif intent == "Get Corrected Version":
        return get_corrected_text(text)  # New function to return corrected text
    else:
        return "Sorry, I couldn't recognize the intent."

def get_corrected_text(text):
    # Regex to find words with the format: original <suggested>
    pattern = r"\b\w+[.,]?\s*<([^>]+)>"

    # Replace the original word with the suggested word inside angle brackets
    corrected_text = re.sub(pattern, r"\1", text)

    return corrected_text

def summarize_text(text, openai_client):
    prompt = f"Please summarize the following text:\n\n{text}\n\nSummary:"

    response = openai_client.chat.completions.create(
        model="gpt-4",
        messages=[
            {
                "role": "user",
                "content": prompt,
            },
        ],
        temperature=0.5,
        max_tokens=150,
    )

    summary = response.choices[0].message.content.strip()
    return summary

def redact_pii(text):
    # Example patterns for PII (you may need to adjust these based on your requirements)
    patterns = {
        r'\b[A-Z][a-z]+ [A-Z][a-z]+\b': '[REDACTED NAME]',  # Names (e.g., John Doe)
        r'\b\d{3}-\d{2}-\d{4}\b': '[REDACTED SSN]',  # Social Security Numbers
        r'\b\w+@\w+\.\w+\b': '[REDACTED EMAIL]',  # Emails
        r'\b\d{3}-\d{3}-\d{4}\b': '[REDACTED PHONE]',  # Phone Numbers
    }

    for pattern, replacement in patterns.items():
        text = re.sub(pattern, replacement, text)

    return text

def extract_entities(text):
    response = text_analytics_client.recognize_entities(documents=[text])[0]

    if not response.is_error:
        entities = [(entity.text, entity.category) for entity in response.entities]
        return entities
    else:
        return "Error in entity extraction."

# Load a sentence transformer model (you can change this to another model if needed)
model = SentenceTransformer('paraphrase-MiniLM-L6-v2')

# Function to compute similarity
def compute_similarity(word1, word2):
    embeddings = model.encode([word1, word2])
    return cosine_similarity([embeddings[0]], [embeddings[1]])[0][0]

def process_word(word, context, openai_client, file_path=None):
    response = None  # Initialize response with None or a default value
    suggested_word = word.content  # Default to the original word if no response

    # If confidence is less than 0.9, predict the word using context
    if word.confidence < 0.9:
        prompt = f"The word '{word.content}' might be incorrect. Suggest a more accurate word, considering it might be slightly distorted or misread. Only suggest if it's reasonably certain, otherwise, just return the original word. Context: {context}"

        # Prepare the request
        messages = [
            {
                "role": "user",
                "content": prompt,
            },
        ]

        # If an image is provided, you would include that in the request
        if file_path:
            messages.append({"role": "system", "content": f"Image: {file_path}"})

        try:
            # Make the API call
            response = openai_client.chat.completions.create(
                model="gpt-4",  # Replace with your Azure OpenAI model deployment name
                messages=messages,
                temperature=0.45,
                max_tokens=100
            )

            # Ensure response is valid before attempting to access choices
            if response and hasattr(response, "choices") and response.choices:
                suggested_word = response.choices[0].message.content.strip()
            else:
                st.error("No valid response received from OpenAI API.")
        except Exception as e:
            st.error(f"Error in processing OpenAI request: {e}")
            response = None  # Ensure response is handled in case of an error

        # Compute similarity
        similarity = compute_similarity(word.content, suggested_word)
        print(f" ...... Similarity between '{word.content}' and '{suggested_word}': {similarity}")

        # Add logic based on similarity
         # If the similarity is low, show both original and suggested word
        if similarity < 0.85:
            return f"{word.content} <{suggested_word}>"  # Show original + suggested word with angle brackets
        else:
            # If the similarity is high, keep the original word
            return word.content
    else:
        return word.content

def analyze_layout(file_path, openai_client):

    # Read the file and analyze it (similar to your original function)
    with open(file_path, 'rb') as file:
        data = file.read()

    # Call Document Intelligence API and analyze the document layout
    document_intelligence_client = DocumentIntelligenceClient(
        endpoint=azure_endpoint, credential=AzureKeyCredential(azure_api_key)
    )
    poller = document_intelligence_client.begin_analyze_document(
        model_id="prebuilt-layout",  # Replace with correct model if necessary
        analyze_request=data,
        content_type="application/octet-stream"
    )
    result: AnalyzeResult = poller.result()

    # Check for handwritten content
    if result.styles and any(style.is_handwritten for style in result.styles):
        print("Document contains handwritten content")
    else:
        print("Document does not contain handwritten content")

    # Check if the document contains text and tables
    has_text = len(result.pages) > 0 and any(len(page.lines) > 0 for page in result.pages)
    has_tables = result.tables is not None and len(result.tables) > 0

    aggregated_text1 = []
    aggregated_text2 = []
    if has_text:
        for page in result.pages:
            aggregated_text1.append(f"Page {page.page_number}:\n")
            page_text = []  # To hold text for the current page
            for line_idx, line in enumerate(page.lines):
                words = get_words(page, line)
                line_text = " ".join(word.content for word in words)
                page_text.append(line_text)

            #aggregated_text1.append("\n".join(page_text) + "\n")
            processed_words = []
            for line in page.lines:
                words = get_words(page, line)
                for word in words:
                    processed_word = process_word(word, "\n".join(page_text), openai_client)
                    processed_words.append(processed_word.strip())
            processed_paragraph = " ".join(processed_words)  # Join words with a space
            aggregated_text1.append(processed_paragraph + " ")
            return " ".join(aggregated_text1)  # Return the combined text as a string
    if has_tables:
        table_output = ""
        for table_idx, table in enumerate(result.tables):
            for cell in table.cells:
                cell_content = cell.content
                processed_words = []
                words = cell_content.split()  # Split the cell content into words
                for word in words:
                    word_obj = type('', (), {'content': word, 'confidence': 0.8})()  # Assuming 0.8 confidence
                    processed_word = process_word(word_obj, cell_content, openai_client)
                    processed_words.append(processed_word)

                processed_cell_content = " ".join(processed_words)
                table_output += f"Row {cell.row_index + 1}, Column {cell.column_index + 1}: {processed_cell_content}\n"
                aggregated_text2.append(processed_cell_content + "\n")
                return " ".join(aggregated_text2)  # Return the combined text as a string

def analyze_document_app():
    st.title("Intelligent Document Processing System (AI Scrawl Wizards)")

    if 'file_path' not in st.session_state:
        st.session_state['file_path'] = None

    # Input fields for Azure OpenAI credentials
    azure_openai_key = st.text_input("Azure OpenAI Key", type="password")
    azure_openai_endpoint = st.text_input("Azure OpenAI Endpoint")

    uploaded_file = st.file_uploader("Upload a file for analysis", type=['jpg', 'png', 'pdf'])

    if uploaded_file:
        file_bytes = uploaded_file.read()
        file_path = "/tmp/uploaded_file." + uploaded_file.name.split('.')[-1]

        with open(file_path, 'wb') as f:
            f.write(file_bytes)

        st.session_state['file_path'] = file_path
        st.success(f"File {uploaded_file.name} uploaded successfully.")

    if st.session_state['file_path'] and azure_openai_key and azure_openai_endpoint:
        if st.button('Run Analysis'):
            # Initialize Azure OpenAI client with user-provided credentials
            openai_client = AzureOpenAI(azure_endpoint=azure_openai_endpoint, api_key=azure_openai_key, api_version="2024-02-15-preview")
            st.write("Running analysis on the uploaded file...")
            result_text = analyze_layout(st.session_state['file_path'], openai_client)
            st.session_state['result_text'] = result_text
            st.session_state['openai_client'] = openai_client  # Save the client in session state

    if 'result_text' in st.session_state:
        st.text_area("Analysis Output", value=st.session_state['result_text'], height=400)

        # Use session state to retain the selected command
        if 'user_command' not in st.session_state:
            st.session_state['user_command'] = "summary"

        # Updated to add the new option "Get Corrected Version"
        st.session_state['user_command'] = st.radio(
            "Select a command:",
            options=["Get Corrected Version", "summary", "RedactPII", "GetEntities"],  # Added "Get Corrected Version"
            index=["Get Corrected Version", "summary", "RedactPII", "GetEntities"].index(st.session_state['user_command']),
            key="user_command_radio"
        )

        if st.button('Run Command'):
            if st.session_state['user_command']:
                openai_client = st.session_state.get('openai_client')  # Retrieve the client from session state
                if openai_client:
                    response_message = process_intent(st.session_state['user_command'], st.session_state['result_text'], openai_client)
                    if st.session_state['user_command'] == "Get Corrected Version":
                        st.text_area("Corrected Text", value=response_message, height=400, key="corrected_text")  # New output area for corrected text
                        # Update the result_text in session state with the corrected version
                        st.session_state['result_text'] = response_message
                    else:
                        st.text_area("Command Output", value=response_message, height=400, key="command_output")
                else:
                    st.error("OpenAI client is not initialized. Please run analysis first.")

            if st.session_state['user_command'] != "Get Corrected Version" and 'result_text' in st.session_state:
                # Perform the selected command on the corrected text if available
                response_message = process_intent(st.session_state['user_command'], st.session_state['result_text'], openai_client)
                st.text_area("Command Output", value=response_message, height=400, key="command_output_final")


if __name__ == "__main__":
    analyze_document_app()
